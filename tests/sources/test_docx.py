from pathlib import Path

import pytest

from oxison.sources.docx import DocxAdapter

docx_mod = pytest.importorskip("docx")


def _make_docx(path: Path) -> None:
    from docx import Document
    d = Document()
    d.add_heading("Spec", level=1)
    d.add_paragraph("requirement one")
    d.add_paragraph("requirement two")
    d.save(str(path))


def test_docx_detect(tmp_path: Path):
    assert DocxAdapter().detect(tmp_path / "spec.docx")
    assert not DocxAdapter().detect(tmp_path / "x.pptx")


def test_docx_extracts_paragraphs(tmp_path: Path):
    f = tmp_path / "spec.docx"
    _make_docx(f)
    res = DocxAdapter().extract(f)
    assert res.status == "ok"
    assert res.unit_count == 1
    u = res.units[0]
    assert u.locator == "docx:spec.docx"
    assert "requirement one" in u.text
    assert "requirement two" in u.text


def test_docx_skips_empty(tmp_path: Path):
    f = tmp_path / "blank.docx"
    from docx import Document
    Document().save(str(f))
    res = DocxAdapter().extract(f)
    assert res.status == "skipped"
    assert res.reason == "no paragraph text"


def test_docx_degrades_when_lib_absent(tmp_path: Path, monkeypatch):
    import builtins
    real_import = builtins.__import__

    def fake_import(name, *args, **kwargs):
        if name == "docx":
            raise ImportError("forced: python-docx absent")
        return real_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", fake_import)
    f = tmp_path / "x.docx"
    f.write_bytes(b"PK fake")
    res = DocxAdapter().extract(f)
    assert res.status == "skipped"
    assert "python-docx not installed" in (res.reason or "")
